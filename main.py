"""
If you are reading this, you probably know what you're doing.
Feel free to read through the code and understand how it works.
"""


# IMPORTS

import os
import math
import docx
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import tkinter as tk
from tkinter import filedialog


# CLASSES

class Colors:
    """
    This class contains ANSI escape codes for text colors and styles.
    """

    RED = "\u001b[31m"
    GREEN = "\u001b[32m"
    YELLOW = "\u001b[33m"
    BLUE = "\u001b[34m"
    MAGENTA = "\u001b[35m"
    CYAN = "\u001b[36m"
    WHITE = "\u001b[37m"
    GRAY = "\u001b[90m"
    RESET = "\u001b[0m"
    BOLD = "\u001b[1m"
    UNDERLINE = "\u001b[4m"

class Value:
    """
    This class represents a numerical value.
    Values exist for easy access to the value and its unit.
    """

    def __init__(self, value: int | float, unit: str):
        self.value = value
        self.unit = unit

    def __str__(self):
        return f"{self.value}{self.unit}"
    
    def __repr__(self):
        return f"{self.value}{self.unit}"
    
    def __eq__(self, other):
        if type(other) == Value:
            return self.value == other.value and self.unit == other.unit
        if type(other) in [int, float]:
            return self.value == other
        return False
    
    def __ne__(self, other):
        return not self.__eq__(other)
    
    def __lt__(self, other):
        if type(other) == Value:
            if self.unit != other.unit:
                raise ValueError("Cannot compare values with different units, with the exception of equality or unitless values.")
            return self.value < other.value
        if type(other) in [int, float]:
            return self.value < other
        return False
    
    def __le__(self, other):
        return self.__lt__(other) or self.__eq__(other)
    
    def __gt__(self, other):
        if type(other) == Value:
            if self.unit != other.unit:
                raise ValueError("Cannot compare values with different units, with the exception of equality or unitless values.")
            return self.value > other.value
        if type(other) in [int, float]:
            return self.value > other
        return False
    
    def __ge__(self, other):
        return self.__gt__(other) or self.__eq__(other)
    
    def __add__(self, other):
        if type(other) == Value:
            if self.unit != other.unit:
                raise ValueError("Cannot add values with different units, with the exception of unitless values.")
            return Value(self.value + other.value, self.unit)
        if type(other) in [int, float]:
            return Value(self.value + other, self.unit)
        return None
    
    def __sub__(self, other):
        if type(other) == Value:
            if self.unit != other.unit:
                raise ValueError("Cannot subtract values with different units, with the exception of unitless values.")
            return Value(self.value - other.value, self.unit)
        if type(other) in [int, float]:
            return Value(self.value - other, self.unit)
        return None
    
    def __mul__(self, other):
        if type(other) == Value:
            if self.unit != other.unit:
                raise ValueError("Cannot multiply values with different units, with the exception of unitless values. This is a limitation of the current implementation.")
            return Value(self.value * other.value, self.unit)
        if type(other) in [int, float]:
            return Value(self.value * other, self.unit)
        return None
    
    def __truediv__(self, other):
        if type(other) == Value:
            if self.unit != other.unit:
                raise ValueError("Cannot divide values with different units, with the exception of unitless values. This is a limitation of the current implementation.")
            return Value(self.value / other.value, self.unit)
        if type(other) in [int, float]:
            return Value(self.value / other, self.unit)
        return None
    
    def __floordiv__(self, other):
        if type(other) == Value:
            if self.unit != other.unit:
                raise ValueError("Cannot divide values with different units, with the exception of unitless values. This is a limitation of the current implementation.")
            return Value(self.value // other.value, self.unit)
        if type(other) in [int, float]:
            return Value(self.value // other, self.unit)
        return None
    
    def __mod__(self, other):
        if type(other) == Value:
            if self.unit != other.unit:
                raise ValueError("Cannot divide values with different units, with the exception of unitless values. This is a limitation of the current implementation.")
            return Value(self.value % other.value, self.unit)
        if type(other) in [int, float]:
            return Value(self.value % other, self.unit)
        return None
    
    def __pow__(self, other):
        if type(other) == Value:
            if self.unit != other.unit:
                raise ValueError("Cannot raise values with different units, with the exception of unitless values. This is a limitation of the current implementation.")
            return Value(self.value ** other.value, self.unit)
        if type(other) in [int, float]:
            return Value(self.value ** other, self.unit)
        return None
    
    def __abs__(self):
        return Value(abs(self.value), self.unit)
    
    def __neg__(self):
        return Value(-self.value, self.unit)
    
    def __pos__(self):
        return Value(self.value, self.unit)
    
    def __int__(self):
        return int(self.value)
    
    def __float__(self):
        return float(self.value)
    
    def __round__(self, n: int = 0):
        return round(self.value, n)
    
    def __ceil__(self):
        return math.ceil(self.value)
    
    def __floor__(self):    
        return math.floor(self.value)

class Range:
    def __init__(self, start: Value, end: Value):
        if start.unit != end.unit:
            raise ValueError("Cannot create a range with different units.")
        # if start > end:
        #     temp = start
        #     start = end
        #     end = temp
        #     del temp
        self.start = start
        self.end = end

    def __str__(self):
        return f"{str(self.start)} to {str(self.end)}"

class Operation:
    """
    This class represents an operation in the process.
    """

    def __init__(self, data: list):
        self.data = data

    def __str__(self) -> str:
        return str(self.data)
    
    def __repr__(self) -> str:
        return str(self)
    
    def get_from(self) -> float:
        return float(self.data[0])
    
    def get_to(self) -> float:
        return float(self.data[1])
    
    def get_duration(self) -> float:
        return self.data[2]
    
    def get_operation(self) -> str:
        return self.data[4]
    
    def get_mud_weight(self) -> Value | None:
        operation = self.get_operation()
        try:
            data = operation.split("@")[1].split(". ")
            for entry in data:
                if "MW" in entry:
                    return Value(float(entry.replace("kg/m³ MW", "").strip()), "kg/m³")
            return None
        except:
            return None
        
class Drilling(Operation):
    """
    This class represents a drilling operation in the process.
    """

    def __str__(self) -> str:
        return str(f"[{convert_time(self.get_from())} - {convert_time(self.get_to())}] Drilling to {self.get_depth()} at {self.get_pump_rate()}. {self.get_mud_weight()} MW. {self.get_ecd()} ECD.")

    def get_depth(self) -> Value | None:
        try:
            operation = self.get_operation()
            return Value(float(operation.split("@")[0].replace("Drilling (to ", "").replace("m)", "")), "m")
        except:
            return None
    
    def get_ecd(self) -> Value | None:
        operation = self.get_operation()
        try:
            data = operation.split("@")[1].split(". ")
            for entry in data:
                if "ECD" in entry:
                    return Value(float(entry.replace("kg/m³ ECD", "").strip()), "kg/m³")
            return None
        except:
            return None
        
    def get_pump_rate(self) -> Value | None:
        operation = self.get_operation()
        try:
            data = operation.split("@")[1].split(". ")
            for entry in data:
                if "m³/min" in entry:
                    return Value(float(entry.replace("m³/min", "").strip()), "m³/min")
            return None
        except:
            return None

class Connection(Operation):
    """
    This class represents a connection operation in the process.
    """

    def __str__(self) -> str:
        return str(f"[{convert_time(self.get_from())} - {convert_time(self.get_to())}] Connection at {self.get_depth()}. {self.get_mud_weight()} MW. {self.get_esd()} ESD. {self.get_static_bp()} BP. {self.get_gas()} B/U.")

    def get_depth(self) -> Value | None:
        try:
            operation = self.get_operation()
            return Value(float(operation.split("@")[1].split(". ")[0].replace("m", "").strip()), "m")
        except:
            return None
    
    def get_esd(self) -> Value | None:
        try:
            operation = self.get_operation()
            data = operation.split("@")[1].split(". ")
            for entry in data:
                if "ESD" in entry:
                    return Value(float(entry.replace("kg/m³ ESD", "").strip()), "kg/m³")
            return None
        except:
            return None
        
    def get_gas(self) -> Value | None:
        try:
            operation = self.get_operation()
            data = operation.split("@")[1].split(". ")
            for entry in data:
                if "KSCM/Day B/U" in entry:
                    return Value(float(entry.replace("KSCM/Day B/U", "").strip()), "KSCM/Day B/U")
                elif "No Gas to Report" in entry:
                    return Value(0, "KSCM/Day")
                elif "Flame B/U" in entry:
                    return Value(0, "KSCM/Day")
            return None
        except Exception as e:
            return None
        
    def get_static_bp(self) -> Value | Range | None:
        try:
            operation = self.get_operation()
            data = operation.split("@")[1].split(". ")
            for entry in data:
                if "BP" in entry:
                    if any(i.lower() in entry.lower() for i in ["No BP", "Closed Choke", "Open Choke"]):
                        return Value(0, "kPa")
                    elif " to " in entry:
                        values = entry.replace("kPa BP", "").split(" to ")
                        start = Value(float(values[0]), "kPA")
                        end = Value(float(values[1]), "kPA")
                        return Range(start, end)
            return None
        except:
            return None
        
class Logger:
    """
    This class represents the logger for the program.
    This is a simply utility class to handle output.
    It is not meant to be used all the time, but makes it easier to output certain messages.
    """

    def error(text: str) -> None:
        """
        Print an error message based on the error code.

        Parameters:
            code (int): The error code to print the message for.

        Returns:
            None
        """

        fprint(f"{Colors.RED}Error: {text}")

    def warn(text: str) -> None:
        """
        Print a warning message.

        Parameters:
            text (str): The warning message to print.

        Returns:
            None
        """

        fprint(f"{Colors.YELLOW}Warning: {text}")


# UTILITY FUNCTIONS

def fprint(*args, end: str = "\n") -> None:
    """
    Print function with built-in color reset for convenience.
    This is to be used instead of the built-in print function because of its automatic color reset.

    Parameters:
        *args (any): Any number of arguments to print, can be any type.

    Returns:
        None
    """

    # Altered print statement to include color reset.
    print(str(*args) + Colors.RESET, end=end)

def clear() -> None:
    """
    Clears the terminal screen using an ANSI escape code.

    Parameters:
        None
    
    Returns:
        None
    """

    # ANSI escape code to clear the screen.
    print("\033[H\033[J", end='')

def pause(newlines: int = 1) -> None:
    """
    Pause the program until the user presses Enter.

    Parameters:
        newlines (int): The number of newlines to print before the prompt. Default is 1.

    Returns:
        None
    """

    # Print newlines and prompt the user to press the enter key.
    fprint(("\n" * newlines) + Colors.BOLD + "Press ENTER to continue.", end="")
    input()

def convert_time(time_object: str | float) -> str | float:
    """
    Converts time from float to string or vice versa.
    Time can be represented as a float between 0 to 1 or a string in the format "HH:MM".

    Parameters:
        time_object (str | float): The time in either format.

    Returns:
        str | float: The time in the other format.
    """

    if type(time_object) == str:
        hour = int(time_object.split(":")[0])
        minute = int(time_object.split(":")[1])
        return float((hour * 60 + minute) / 1440)
    elif type(time_object) == float:
        hour = int(time_object * 1440 / 60)
        minute = int(time_object * 1440 % 60)
        return f"{hour}:{minute:02d}"
    else:
        return None
    
def validate_time(timestring: str) -> bool:
    """
    Validate a time string in the format "HH:MM" in 24 hour time.
    This function will check if the time string is valid and return a boolean.

    Parameters:
        timestring (str): The time string to validate.

    Returns:
        bool: True if the time string is valid, False otherwise.
    """

    try:
        hour = int(timestring.split(":")[0])
        minute = int(timestring.split(":")[1])
        if hour < 0 or hour > 23:
            return False
        if minute < 0 or minute > 59:
            return False
        return True
    except:
        return False


def parse_sheet(data: pd.DataFrame) -> list:
    """
    Parse the data from a sheet in the Excel file.
    This function will parse the data from the sheet and return a list of operations.

    Parameters:
        data (pd.DataFrame): The data from the sheet to parse.

    Returns:
        list: A list of operations parsed from the data.
    """

    # Convert the data to a list of lists.
    data_list = data.values.tolist()

    # Initialize a new list to store the parsed data.
    new = []

    # Iterate over the data and parse the operations.
    for index, row in enumerate(data_list):
        if index < 21:
            continue
        sub = [j for j in row if type(j) != float]
        new.append(sub)

    # Validate the parsed data.
    operations = []
    for index, row in enumerate(new):
        if index == 0 and row[0] != "0":
            Logger.error("Operation summary is malformed, please check the Excel file and try again.")
            Logger.error("This is most likely because the first operation does not start at 0:00.")
            return
        
        if len(row) <= 1:
            break

        if len(row) < 5:
            continue

        try:
            connection = int(row[3])
            if connection > 0:
                operations.append(Connection(row))
                continue
        except ValueError:
            pass

        if row[3] == "C":
            operations.append(Connection(row))
        elif row[3] == "D":
            operations.append(Drilling(row))

    # Return the parsed data.
    return operations

def get_range_depth(data: list) -> Range:
    """
    Get the range of depths from a list of a list of operations.

    Parameters:
        data (list): The 2D list of operations to get the range from. This is a list of days, each containing a list of operations.

    Returns:
        Range: The range of depths.
    """

    # Initialize the depth list.
    depths = []

    # Iterate over the operations and extract the depths.
    for day in data:
            for operation in day:
                if operation.get_depth() != None and type(operation.get_depth()) == Value:
                    depths.append(operation.get_depth().value)

    # Calculate the minimum and maximum depths.
    depth_minimum = min(depths)
    depth_maximum = max(depths)

    # Return the range of depths.
    return Range(Value(depth_minimum, "m"), Value(depth_maximum, "m"))


# MAIN FUNCTION

def main() -> None:
    # Clear the screen and print the welcome message.
    clear()
    fprint(f"{Colors.BOLD}Welcome to the Drilling Report Parser!")
    pause()
    clear()

    # Open a file dialog to select the Excel file.
    fprint("Opening file dialog to select the Drilling Operations Report Excel file...")
    fprint("Any messages printed after this are for debugging purposes and can be ignored.\n")
    root = tk.Tk()
    root.withdraw()
    filename = filedialog.askopenfilename(
        initialdir=os.getcwd(),
        title="Select the Drilling Operations Report Excel file",
        filetypes=(
            ("Excel file", "*.xlsb *.xls *.xlsx"),
            ("All files (may encounter errors)", "*.*")
        )
    )

    # Load the Excel file and get the sheet names.
    fprint(f"Loading Excel file '{filename}'...")
    try:
        data = pd.ExcelFile(filename)
    except Exception as e:
        Logger.error(f"An error occurred while loading the Excel file: {e}")
        pause()
        return
    
    pause()

    # Get sheet names, and filter only the sheets with dates matching the format (example: "2024 Dec-18").
    sheet_names = data.sheet_names
    sheets = []
    for i in sheet_names:
        try:
            datetime.strptime(i, '%Y %b-%d')
            sheets.append(i)
        except ValueError:
            continue
        
    # Check if there are any sheets with dates.
    if len(sheets) == 0:
        # Print error message and return if no sheets are found.
        Logger.error("No DORs found. Please check the Excel file and try again.")
        return
    
    # Ask the user for a date range.
    while True:
        start_sheet = None
        start_time = None
        end_sheet = None
        end_time = None

        # Get the start date.
        while True:
            clear()
            fprint(f"{Colors.BOLD}Please input the number next to the day you'd like to start at:")
            for index, sheet in enumerate(sheets):
                fprint(f"\t{Colors.BOLD}{index + 1}. {Colors.RESET}{sheet}")
            try:
                option = int(input("> "))
                if option < 1 or option > len(sheets):
                    raise ValueError
                start_sheet = sheets[option - 1]
                break
            except ValueError:
                Logger.error("Invalid input. Please enter a valid number.")
                pause()

        # Get the start time.
        while True:
            clear()
            fprint(f"{Colors.BOLD}Selected start date: {Colors.RESET}{start_sheet}")
            fprint(f"{Colors.BOLD}Please input the 24-hour time you'd like to start at (HH:MM):")
            try:
                start_time = input("> ")
                if start_time == "":
                    start_time = "00:00"
                if not validate_time(start_time):
                    raise ValueError
                break
            except ValueError:
                Logger.error("Invalid input. Please enter the time in the 24-hour time format 'HH:MM'. You can leave this blank to start at 00:00.")
                pause()

        # Get the end date.
        while True:
            clear()
            fprint(f"{Colors.BOLD}Selected start date: {Colors.RESET}{start_sheet}")
            fprint(f"{Colors.BOLD}Selected start time: {Colors.RESET}{start_time}")
            fprint(f"{Colors.BOLD}Please input the number next to the day you'd like to end at:")
            for index, sheet in enumerate(sheets):
                fprint(f"\t{index + 1}. {sheet}")
            try:
                option = int(input("> "))
                if option < 1 or option > len(sheets):
                    raise ValueError
                end_sheet = sheets[option - 1]
                break
            except ValueError:
                Logger.error("Invalid input. Please enter a valid number.")
                pause()

        # Get the end time.
        while True:
            clear()
            fprint(f"{Colors.BOLD}Selected start date: {Colors.RESET}{start_sheet}")
            fprint(f"{Colors.BOLD}Selected start time: {Colors.RESET}{start_time}")
            fprint(f"{Colors.BOLD}Selected end date: {Colors.RESET}{end_sheet}")
            fprint(f"{Colors.BOLD}Please input the 24-hour time you'd like to end at (HH:MM):")
            try:
                end_time = input("> ")
                if end_time == "":
                    end_time = "23:59"
                if not validate_time(end_time):
                    raise ValueError
                break
            except ValueError:
                Logger.error("Invalid input. Please enter the time in the 24-hour time format 'HH:MM'. You can leave this blank to end at 23:59.")
                pause()

        # Validate the date range.
        # Check if the start and end dates are in the wrong order.
        if sheets.index(start_sheet) > sheets.index(end_sheet):
            Logger.error("Invalid date range. Please select an end date after the start date.")
            pause()
            continue

        # If the start and end dates are the same, check if the start time is after the end time.
        if start_sheet == end_sheet:
            if convert_time(start_time) > convert_time(end_time):
                Logger.error("Invalid time range. Please select an end time after the start time.")
                pause()
                continue

        # If the date range is valid, break the loop.
        break

    # Print the selected date range.x
    clear()
    fprint(f"{Colors.BOLD}Selected start date: {Colors.RESET}{start_sheet}")
    fprint(f"{Colors.BOLD}Selected start time: {Colors.RESET}{start_time}")
    fprint(f"{Colors.BOLD}Selected end date: {Colors.RESET}{end_sheet}")
    fprint(f"{Colors.BOLD}Selected end time: {Colors.RESET}{end_time}")
    fprint()

    # Get the data from the selected sheets.
    start_index = sheets.index(start_sheet)
    end_index = sheets.index(end_sheet)
    days = []
    for i in range(start_index, end_index + 1):
        days.append(parse_sheet(data.parse(sheets[i], dtype=str)))

    # Filter the data based on the start and end times.
    start_time = convert_time(start_time)
    if end_time == "23:59":
        end_time = 1
    else:
        end_time = convert_time(end_time)

    for operation in days[0].copy():
        if operation.get_from() < start_time:
            days[0].remove(operation)

    for operation in days[-1].copy():
        if operation.get_to() > end_time:
            days[-1].remove(operation)

    # Get the calculated data.
    rows = []
    last_depth = None
    for index, day in enumerate(days, start=1):
        # Initialize the data lists.
        depths = []
        mud_weights = []
        pump_rates = []
        static_bps = []
        ecds = []
        esds = []
        gases = []

        # Iterate over the operations and extract the data.
        for operation in day:
            depths.append(operation.get_depth())
            mud_weights.append(operation.get_mud_weight())
            if type(operation) == Drilling:
                pump_rates.append(operation.get_pump_rate())
                ecds.append(operation.get_ecd())
            elif type(operation) == Connection:
                esds.append(operation.get_esd())
                gases.append(operation.get_gas())
                static_bp = operation.get_static_bp()
                if type(static_bp) == Value:
                    static_bps.append(static_bp)
                elif type(static_bp) == Range:
                    static_bps.append(static_bp.start)
                    static_bps.append(static_bp.end)

        # Calcualte minimum and maximum values for depth.
        try:
            clean_depths = [i.value for i in depths if i != None]
            if len(clean_depths) < len(depths):
                Logger.warn(f"Depths data incomplete for Day #{index}.")
            depth_minimum = min(clean_depths) if last_depth == None else last_depth
            depth_maximum = max(clean_depths)
            depth = f"{depth_minimum:.0f} – {depth_maximum:.0f}"
            if depth_minimum == depth_maximum:
                depth = f"{depth_minimum:.0f}"
            last_depth = depth_maximum
        except:
            Logger.warn(f"No depth reported for Day #{index}.")
            depth = "No depth reported"

        # Calculate minimum and maximum values for mud weight.
        try:
            clean_mud_weights = [i.value for i in mud_weights if i != None]
            if len(clean_mud_weights) < len(mud_weights):
                Logger.warn(f"Mud weight data incomplete for Day #{index}.")
            mud_weight_minimum = min(clean_mud_weights)
            mud_weight_maximum = max(clean_mud_weights)
            mud_weight = f"{mud_weight_minimum:.0f} – {mud_weight_maximum:.0f}"
            if mud_weight_minimum == mud_weight_maximum:
                mud_weight = f"{mud_weight_minimum:.0f}"
        except:
            Logger.warn(f"No mud weight reported for Day #{index}.")
            mud_weight = "No mud weight reported"

        # Calculate minimum and maximum values for pump rate.
        try:
            clean_pump_rates = [i.value for i in pump_rates if i != None]
            if len(clean_pump_rates) < len(pump_rates):
                Logger.warn(f"Pump rate data incomplete for Day #{index}.")
            pump_rate_minimum = min(clean_pump_rates)
            pump_rate_maximum = max(clean_pump_rates)
            pump_rate = f"{round(pump_rate_minimum, 2)} – {round(pump_rate_maximum, 2)}"
            if pump_rate_minimum == pump_rate_maximum:
                pump_rate = f"{round(pump_rate_minimum, 2)}"
        except:
            Logger.warn(f"No pump rate reported for Day #{index}.")
            pump_rate = "No pump rate reported"

        # Dynamic BP is always "Line Restriction". Can also be changed later if needed.
        dynamic_bp = "Line Restriction"

        # Calculate minimum and maximum values for static BP.
        try:
            clean_static_bps = [i.value for i in static_bps if i != None]
            if len(clean_static_bps) < len(static_bps):
                Logger.warn(f"Static BP data incomplete for Day #{index}.")
            static_bp_minimum = min(clean_static_bps)
            static_bp_maximum = max(clean_static_bps)

            if static_bp_minimum == 0 and static_bp_maximum != 0:
                static_bp_minimum = min([i for i in clean_static_bps if i != 0])

            static_bp = f"{static_bp_minimum:.0f} – {static_bp_maximum:.0f}"
            if static_bp_minimum == 0 and static_bp_maximum == 0:
                static_bp = "No BP"
            elif static_bp_minimum == static_bp_maximum:
                static_bp = f"{static_bp_minimum:.0f}"
        except:
            Logger.warn(f"No static BP reported for Day #{index}.")
            static_bp = "No static BP reported"
        
        # Calculate average ECD and ESD.
        try:
            clean_ecds = [i.value for i in ecds if i != None]
            clean_esds = [i.value for i in esds if i != None]
            if len(clean_ecds) < len(ecds):
                Logger.warn(f"ECD data incomplete for Day #{index}.")
            if len(clean_esds) < len(esds):
                Logger.warn(f"ESD data incomplete for Day #{index}.")
            ecd_average = round(sum(clean_ecds) / len(clean_ecds)) if len(clean_ecds) > 0 else "--" 
            esd_average = round(sum(clean_esds) / len(clean_esds)) if len(clean_esds) > 0 else "--"
            ecd_esd_average = f"{ecd_average}/{esd_average}"
        except:
            Logger.warn(f"No ESD/ECD reported for Day #{index}.")
            ecd_esd_average = "No ESD/ECD reported"

        # Calculate maximum value for gas.
        try:
            clean_gases = [i.value for i in gases if i != None]
            if len(clean_gases) < len(gases):
                Logger.warn(f"Gas data incomplete for Day #{index}.")
            gas_maximum = max(clean_gases)
            if gas_maximum == 0:
                gas = "No B/U gas reported"
            else:
                gas = f"Max {gas_maximum:.1f} kscm/d B/U gas reported"
        except Exception:
            Logger.warn(f"No gas reported for Day #{index}.")
            gas = "No B/U gas reported"

        row = [depth, mud_weight, pump_rate, dynamic_bp, static_bp, ecd_esd_average, gas]
        rows.append(row)

        
    fprint(f"\n{Colors.BOLD}If you see any warnings, you may ignore them if the data is not reported in the DOR.")
    pause()

    # Format the date range.
    start_date = datetime.strptime(start_sheet, '%Y %b-%d')
    end_date = datetime.strptime(end_sheet, '%Y %b-%d')
    if start_date.month == end_date.month and start_date.year == end_date.year:
        date_range = f"{start_date.strftime('%b')} {start_date.day} - {end_date.day}, {start_date.year}"
    elif start_date.year == end_date.year:
        date_range = f"{start_date.strftime('%b')} {start_date.day} - {end_date.strftime('%b')} {end_date.day}, {start_date.year}"
    else:
        date_range = f"{start_date.strftime('%b')} {start_date.day}, {start_date.year} - {end_date.strftime('%b')} {end_date.day}, {end_date.year}"
    
    # Fetch maximum depth for the range.
    maximum_depth = get_range_depth(days).end

    # Create a Word document to write the report to.
    document = Document()
    document.sections[0].left_margin = Cm(1.27)
    document.sections[0].right_margin = Cm(1.27)
    document.sections[0].top_margin = Cm(1.27)
    document.sections[0].bottom_margin = Cm(1.27)
    date_range_paragraph = document.add_paragraph(f"{date_range} | Drilling to {maximum_depth.value:.0f} {maximum_depth.unit}MD")
    run = date_range_paragraph.runs[0]
    run.font.size = Pt(9)
    run.font.name = 'Clear Sans'
    table = document.add_table(rows=1 + len(days), cols=7)
    
    table_style = """
    <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>
    </w:tblBorders>
    """

    # Apply border style to the table.
    table_borders = parse_xml(table_style)
    table._tblPr.addnext(table_borders)


    header_cells = table.rows[0].cells
    headers = [
        "Drilling Interval (mMD)",
        "Mud Weight (kg/m³)",
        "Pump Rate (m³/min)",
        "Dynamic BP (kPa)",
        "Static BP (kPa)",
        "Avg ECD/ESD (kg/m³)",
        "Comments"
    ]

    for index, header in enumerate(header_cells):
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = headers[index]
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        header_paragraph.runs[0].font.size = Pt(10)
        header_paragraph.runs[0].font.name = 'Oswald'
        header_paragraph.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        element = parse_xml(r'<w:shd {} w:fill="2D2F3E"/>'.format(nsdecls('w')))
        header._tc.get_or_add_tcPr().append(element)

    # Set row data.
    column_widths = [2.72, 2.22, 2, 2.25, 2.25, 2.5, 5.11]
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, value in enumerate(row):
            row_paragraph = cells[j].paragraphs[0]
            row_paragraph.text = value
            row_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row_paragraph.runs[0].font.size = Pt(9)
            row_paragraph.runs[0].font.name = 'Clear Sans'
            cells[j].width = Cm(column_widths[j])
            if i % 2 == 0:
                element = parse_xml(r'<w:shd {} w:fill="CCCCCC"/>'.format(nsdecls('w')))
                cells[j]._tc.get_or_add_tcPr().append(element)  
            else:
                element = parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.format(nsdecls('w')))
                cells[j]._tc.get_or_add_tcPr().append(element)

    # Print success message and open the file dialog to save the report.
    while True:
        try:
            clear()
            fprint(f"\n{Colors.BOLD}Report generated successfully!")
            fprint("Opening a file dialog to save the report...")
            fprint("Any messages printed after this are for debugging purposes and can be ignored.\n")

            filename = filedialog.asksaveasfilename(
                initialdir=os.getcwd(),
                title="Save the report as",
                filetypes=(
                    ("Word document", "*.docx"),
                    ("All files", "*.*")
                ),
                defaultextension=".docx"
            )
            if filename == "":
                Logger.error("No filename entered. Try again.")
                pause()
                continue
            if not filename.endswith(".docx"):
                filename += ".docx"

            document.save(filename)

            # Print success message and exit the program.
            fprint(f"\n{Colors.BOLD}Report saved successfully!")
            fprint(f"Report saved as {filename}.")
            fprint(f"What would you like to do next?")
            fprint(f"\t1. Exit the program.")
            fprint(f"\t2. Generate another report.")
            while True:
                try:
                    option = int(input("> "))
                    if option == 1:
                        clear()
                        fprint("Goodbye! Thank you for using the Drilling Report Parser.")
                        return
                    elif option == 2:
                        main()
                    else:
                        raise ValueError
                except ValueError:
                    Logger.error("Invalid input. Please enter a valid number.")
                    pause()
        except Exception as e:
            Logger.error(f"An error occurred while saving the report. Please try again.")
            Logger.error(f"Ensure that the filename is valid and the file is not open in another program if it already exists.")
            pause()


if __name__ == "__main__":
    main()